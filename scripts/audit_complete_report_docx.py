"""Audit QRFL_Complete_Report.docx against all diagrams and measured tables."""
from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document

ROOT = Path(r"C:\Quantum Pqc RFHL\qrfl-artifacts")
DOCX = Path(r"C:\Quantum Pqc RFHL\docs\QRFL_Complete_Report.docx")
RES = ROOT / "results"
FIG = ROOT / "figures" / "output"
FIGS = ROOT.parent / "figs"
TMP = ROOT.parent / "docs" / "_pdf_assets"


def main() -> None:
    doc = Document(str(DOCX))

    with zipfile.ZipFile(DOCX) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
        print(f"DOCX media files: {len(media)}")
        for n in media:
            info = z.getinfo(n)
            print(f"  {n}  {info.file_size/1024:.1f} KB")

    drawing_count = sum(
        1 for p in doc.paragraphs if p._element.xpath('.//*[local-name()="drawing"]')
    )
    print(f"Paragraphs with drawings: {drawing_count}")
    print(f"Tables in DOCX: {len(doc.tables)}")

    print("\n=== DOCX HEADINGS ===")
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            print(f"  [{p.style.name}] {p.text[:110]}")

    src_pngs = sorted(list(FIG.glob("*.png")) + list(TMP.glob("chart_*.png")))
    arch = list(FIGS.glob("fig2*.png"))
    mod = list(FIGS.glob("modality*.png"))
    print("\n=== SOURCE DIAGRAMS ON DISK ===")
    for p in src_pngs + arch + mod:
        try:
            rel = p.relative_to(ROOT.parent)
        except ValueError:
            rel = p
        print(f"  {rel}  {p.stat().st_size/1024:.1f}KB")

    print("\n=== DOCX TABLES (header + row count) ===")
    for i, t in enumerate(doc.tables, 1):
        hdr = [c.text.strip().replace("\n", " ")[:45] for c in t.rows[0].cells]
        nrows = len(t.rows)
        ncols = len(t.columns)
        sample = ""
        if nrows > 1:
            sample = " | ".join(c.text.strip()[:28] for c in t.rows[1].cells)
        def _safe(s: str) -> str:
            return s.encode("ascii", "replace").decode("ascii")

        print(_safe(f"T{i:02d} ({nrows}x{ncols}): {hdr}"))
        print(_safe(f"     row1: {sample[:140]}"))

    blob = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                blob += "\n" + c.text

    macros = json.loads((RES / "macros.json").read_text(encoding="utf-8"))
    pqc = pd.read_csv(RES / "pqc" / "summary.csv")
    fl = pd.read_csv(RES / "fl" / "all_results.csv")
    stats = pd.read_csv(RES / "statistics" / "fl_statistical_tests.csv")
    sc = pd.read_csv(RES / "forecasting" / "scenarios.csv")
    res = pd.read_csv(RES / "resource" / "summary.csv")
    docker = json.loads((RES / "blockchain" / "docker_stats.json").read_text(encoding="utf-8"))

    print("\n=== VALUE PRESENCE CHECKS ===")
    must = {
        "FL seeds 10": macros["FLNumSeeds"],
        "FL rounds 50": macros["FLNumRounds"],
        "IID accuracy 65.3": macros["FLIIDAccuracyClassical"],
        "AUROC 0.566": macros["FLIIDAUROC"],
        "Bootstrap lower 2042": macros["ForecastBootstrapCILower"][:4],
        "ML-KEM-768 encaps 0.128": "0.128",
        "ML-DSA-65 sign 0.850": "0.850",
        "SLH sign 536": "536",
        "HLF classical 80.8": "80.8",
        "HLF native 84.2": "84.2",
        "Holdout MAE 563": "563",
        "Shapiro present": "Shapiro",
        "equivalent present": "equivalent",
        "Docker hospitala": "hospitala",
        "Ablation Balanced": "Balanced",
        "label_flip present": "label_flip",
        "Resource encaps op": "ML-KEM-768_encaps",
        "Determinism passed": macros["FLDeterminismCheck"],
    }
    fails = []
    for label, needle in must.items():
        ok = str(needle) in blob
        print(f"  {'PASS' if ok else 'FAIL'}  {label}")
        if not ok:
            fails.append(label)

    print("\n=== PQC TABLE COVERAGE ===")
    missing_pqc = []
    for r in pqc.itertuples():
        found = False
        for t in doc.tables:
            for row in t.rows[1:]:
                cells = [c.text for c in row.cells]
                if len(cells) >= 2 and r.scheme == cells[0] and r.operation == cells[1]:
                    found = True
                    break
            if found:
                break
        if not found:
            missing_pqc.append(f"{r.scheme}/{r.operation}")
    print(f"  CSV ops={len(pqc)} missing={len(missing_pqc)}")
    for m in missing_pqc:
        print(f"    MISS {m}")

    print("\n=== STATS TABLE COVERAGE ===")
    miss_stats = []
    for r in stats.itertuples():
        ok = False
        for t in doc.tables:
            for row in t.rows[1:]:
                cells = [c.text for c in row.cells]
                joined = " ".join(cells)
                if r.metric in cells and r.mode_a in joined and r.mode_b in joined:
                    ok = True
                    break
            if ok:
                break
        if not ok:
            miss_stats.append(f"{r.metric} {r.mode_a} vs {r.mode_b}")
    print(f"  CSV rows={len(stats)} missing={len(miss_stats)}")
    for m in miss_stats:
        print(f"    MISS {m}")

    print("\n=== SCENARIO COVERAGE ===")
    for r in sc.itertuples():
        ok = r.scenario in blob
        print(f"  {'PASS' if ok else 'FAIL'}  {r.scenario} η={int(r.eta)} lower={r.ci_lower_95:.2f}")

    print("\n=== RESOURCE COVERAGE ===")
    for r in res.itertuples():
        print(f"  {'PASS' if r.operation in blob else 'FAIL'}  {r.operation}")

    print("\n=== DOCKER COVERAGE ===")
    for d in docker:
        print(f"  {'PASS' if d['Name'] in blob else 'FAIL'}  {d['Name']}")

    print("\n=== FL IID ACCURACY STRINGS ===")
    iid = fl[
        (fl.alpha == 1)
        & (fl.aggregator == "fedavg")
        & (fl.attack == "none")
        & (fl.malicious_fraction == 0)
    ]
    for c in sorted(iid.num_clients.unique()):
        acc = iid[iid.num_clients == c].accuracy.mean() * 100
        needle = f"{acc:.1f}%"
        print(f"  clients={int(c)} {needle}  {'PASS' if needle in blob else 'FAIL'}")

    byz = fl[(fl.security_mode == "native_pq") & (fl.num_clients == 25)]
    bg = byz.groupby(["attack", "malicious_fraction", "aggregator"]).size()
    byz_rows = 0
    for t in doc.tables:
        hdr = [c.text for c in t.rows[0].cells]
        if any("Attack" in h for h in hdr) and any(("Agg" in h) or ("Aggregator" in h) for h in hdr):
            byz_rows = len(t.rows) - 1
    print(f"\n=== BYZANTINE groups CSV={len(bg)} DOCX rows={byz_rows} ===")

    expected_figs = [
        "fl_overhead_results.png",
        "forecasting_model_comparison.png",
        "hlf_phase_latencies.png",
        "pqc_overhead_results.png",
        "sensitivity_analysis.png",
        "threshold_crossing_uncertainty.png",
        "chart_fl_latency_modes.png",
        "chart_fl_accuracy.png",
        "chart_hlf_phases.png",
        "chart_pqc_ops.png",
        "chart_pqc_all_ops.png",
        "chart_ablation.png",
        "chart_scenarios.png",
        "chart_resource.png",
        "chart_bc_summary.png",
        "chart_byzantine.png",
    ]
    print("\n=== DIAGRAM FILE EXISTENCE ===")
    miss_fig = []
    for name in expected_figs:
        exists = (FIG / name).exists() or (TMP / name).exists()
        print(f"  {'OK' if exists else 'MISS'}  {name}")
        if not exists:
            miss_fig.append(name)

    # Inventory section claims vs reality
    inv_ok = "fl_overhead_results.png" in blob and "chart_byzantine.png" in blob
    print(f"\nFigure inventory lists files in text: {'PASS' if inv_ok else 'FAIL'}")

    # Count expected embedded images:
    # 6 pipeline + 4 make_extra_charts + 6 make_more_charts + arch + optional modality
    expected_embedded = 6 + 4 + 6  # 16 base charts
    if arch:
        expected_embedded += 1
    if mod:
        expected_embedded += 1
    print("\n=== SUMMARY ===")
    print(f"DOCX images (media): {len(media)}  (expected ~{expected_embedded})")
    print(f"DOCX tables: {len(doc.tables)}")
    print(f"Pipeline PNGs: {len(list(FIG.glob('*.png')))}")
    print(f"Generated chart PNGs: {len(list(TMP.glob('chart_*.png')))}")
    print(f"Value check fails: {len(fails)} -> {fails}")
    print(f"PQC missing ops: {len(missing_pqc)}")
    print(f"Stats missing: {len(miss_stats)}")
    print(f"Diagram files missing on disk: {miss_fig}")
    print(f"Byzantine row match: {byz_rows == len(bg)}")
    print(f"Image count match: {len(media) >= 16}")


if __name__ == "__main__":
    main()
